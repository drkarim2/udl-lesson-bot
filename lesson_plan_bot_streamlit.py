
import streamlit as st
from docx import Document

def generate_lesson_plan(data):
    doc = Document()
    doc.add_heading('خطة درس قائمة على التصميم الشامل', 0)

    doc.add_heading('📚 موضوع الدرس:', level=1)
    doc.add_paragraph(data["topic"])

    doc.add_heading('🎓 الصف الدراسي:', level=1)
    doc.add_paragraph(data["grade"])

    doc.add_heading('🎯 الأهداف التعليمية:', level=1)
    doc.add_paragraph(data["objectives"])

    doc.add_heading('👨‍🏫 أنماط التعلم المستهدفة:', level=1)
    doc.add_paragraph(data["learning_styles"])

    doc.add_heading('♿ احتياجات خاصة:', level=1)
    doc.add_paragraph("نعم" if data["has_special_needs"] else "لا")

    doc.add_heading('📋 الأنشطة المقترحة:', level=1)
    doc.add_paragraph("- فيديو تعليمي مبسط للمتعلمين البصريين\n"
                      "- تسجيل صوتي يشرح المفاهيم للمتعلمين السمعيين\n"
                      "- نشاط عملي بسيط باستخدام ورق وقلم للمتعلمين الحركيين")

    doc.add_heading('📊 أساليب التقييم:', level=1)
    doc.add_paragraph("- أسئلة تفاعلية عبر Kahoot\n"
                      "- ورقة عمل مصورة\n"
                      "- نقاش جماعي شفوي")

    filename = f"lesson_plan_{data['topic'].replace(' ', '_')}.docx"
    doc.save(filename)
    return filename

st.title("🤖 البوت الذكي لإعداد خطة درس شاملة")
st.markdown("أدخل بيانات الدرس وسنقوم بإنشاء خطة جاهزة للطباعة بناءً على مبادئ التصميم الشامل (UDL)")

topic = st.text_input("1️⃣ ما موضوع الدرس؟")
grade = st.text_input("2️⃣ ما الصف الدراسي؟")
objectives = st.text_area("3️⃣ ما الأهداف التعليمية؟")
has_special_needs = st.radio("4️⃣ هل يوجد طلاب من ذوي الاحتياجات الخاصة؟", ["نعم", "لا"])
learning_styles = st.multiselect("5️⃣ ما أنماط التعلم التي تريد مراعاتها؟", ["سمعي", "بصري", "حركي"])

if st.button("🔧 إنشاء خطة الدرس"):
    if topic and grade and objectives and learning_styles:
        data = {
            "topic": topic,
            "grade": grade,
            "objectives": objectives,
            "has_special_needs": has_special_needs == "نعم",
            "learning_styles": "، ".join(learning_styles)
        }
        file_name = generate_lesson_plan(data)
        with open(file_name, "rb") as f:
            st.success("✅ تم إنشاء خطة الدرس بنجاح!")
            st.download_button("📥 تحميل الخطة كملف Word", f, file_name)
    else:
        st.warning("⚠️ الرجاء ملء جميع الحقول المطلوبة قبل إنشاء الخطة.")
